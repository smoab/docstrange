"""DOCX file processor."""

import os
from typing import Dict, Any

from .base import BaseProcessor
from ..result import ConversionResult
from ..exceptions import ConversionError, FileNotFoundError


class DOCXProcessor(BaseProcessor):
    """Processor for Microsoft Word DOCX files."""
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if this processor can handle the file
        """
        if not os.path.exists(file_path):
            return False
        
        # Check file extension
        _, ext = os.path.splitext(file_path.lower())
        return ext in ['.docx', '.doc']
    
    def process(self, file_path: str) -> ConversionResult:
        """Process the DOCX file and return a conversion result.
        
        Args:
            file_path: Path to the DOCX file to process
            
        Returns:
            ConversionResult containing the processed content
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ConversionError: If processing fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            from docx import Document
            
            metadata = self.get_metadata(file_path)
            content_parts = []
            
            doc = Document(file_path)
            
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections)
            })
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if this is a heading
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            content_parts.append(f"\n{'#' * min(level_num, 6)} {paragraph.text}\n")
                        except ValueError:
                            content_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                if self.preserve_layout:
                    content_parts.append("\n### Table\n")
                
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_content.append(' | '.join(row_content))
                
                if table_content:
                    # Add header row separator
                    if len(table_content) > 1:
                        header_separator = ' | '.join(['---'] * len(table_content[0].split(' | ')))
                        table_content.insert(1, header_separator)
                    
                    content_parts.append('\n'.join(table_content))
                    content_parts.append('\n')
            
            content = '\n'.join(content_parts)
            
            # Clean up the content
            content = self._clean_content(content)
            
            return ConversionResult(content, metadata)
            
        except ImportError:
            raise ConversionError("python-docx is required for DOCX processing. Install it with: pip install python-docx")
        except Exception as e:
            if isinstance(e, (FileNotFoundError, ConversionError)):
                raise
            raise ConversionError(f"Failed to process DOCX file {file_path}: {str(e)}")
    
    def _clean_content(self, content: str) -> str:
        """Clean up the extracted DOCX content.
        
        Args:
            content: Raw DOCX text content
            
        Returns:
            Cleaned text content
        """
        # Remove excessive whitespace and normalize
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive whitespace
            line = ' '.join(line.split())
            if line.strip():
                cleaned_lines.append(line)
        
        # Join lines and add proper spacing
        content = '\n'.join(cleaned_lines)
        
        # Add spacing around headers
        content = content.replace('## ', '\n## ')
        content = content.replace('### ', '\n### ')
        
        return content.strip() 